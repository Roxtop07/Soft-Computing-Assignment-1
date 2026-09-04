"""
Builds a beautifully styled Word Document (.docx) for the Assignment 1 Case Study Report.
Includes title, metadata, formatted tables, headers, and embeds high-resolution plots.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

MD_PATH = os.path.join(os.path.dirname(__file__), "Case_Study_Report_Intelligent_AC.md")
DOCX_PATH = os.path.join(os.path.dirname(__file__), "Case_Study_Report_Intelligent_AC.docx")
FIGURES_DIR = os.path.join(os.path.dirname(__file__), "figures")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_docx():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Document Header / Course Banner
    p_course = doc.add_paragraph()
    r_course = p_course.add_run("COURSE: SOFT COMPUTING (STDA2102) | ASSIGNMENT 1: CASE STUDY REPORT")
    r_course.font.name = "Arial"
    r_course.font.size = Pt(10)
    r_course.font.bold = True
    r_course.font.color.rgb = RGBColor(14, 165, 233)
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Design of an Intelligent Air Conditioning System\nUsing Fuzzy Logic")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(16)

    # Metadata Box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    cells = [meta_table.cell(0, 0), meta_table.cell(0, 1), meta_table.cell(1, 0), meta_table.cell(1, 1)]
    texts = [
        ("Course Code & Title:", "STDA2102 - Soft Computing (EL1)"),
        ("Evaluation Weightage:", "50 Marks (10% Total CCE)"),
        ("Student Author:", "Manish Kumar"),
        ("Academic Term:", "Fall / Spring Term 2026")
    ]
    for cell, (lbl, val) in zip(cells, texts):
        cell.width = Inches(3.2)
        set_cell_background(cell, "F1F5F9")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{lbl} ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(val)
        r2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Executive Summary / Abstract
    h1 = doc.add_heading("Executive Summary & Abstract", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.paragraph_format.space_after = Pt(10)
    p_abs.add_run(
        "Heating, Ventilation, and Air Conditioning (HVAC) systems constitute over 50% of the aggregate electrical energy "
        "consumed in modern residential and commercial architectures. Conventional climate control methodologies rely predominantly "
        "on classical thermostats operating under a bivalent 'Bang-Bang' on-off regime or linear Single-Input Single-Output (SISO) PID controllers. "
        "Bang-bang thermostats impose mechanical shock on compressor windings, cause severe cyclical temperature overshoots (±1.5°C to ±2.5°C), "
        "and introduce pronounced thermal discomfort. Conversely, while PID controllers offer continuous feedback, they require exact linear mathematical "
        "modeling of building thermodynamic state variables, which are non-linear, time-variant, and heavily influenced by relative humidity.\n\n"
        "This case study investigates the design, mathematical formulation, simulation, and hardware implementation of an Intelligent Multi-Variable "
        "Fuzzy Logic Controller (FLC). Operating on fuzzy set theory (Zadeh, 1965) and linguistic inference (Mamdani, 1975), the proposed controller "
        "integrates ambient room temperature (°C) and relative humidity (%) to modulate the continuous rotational speed (0%–100%) of a variable-frequency "
        "inverter compressor and indoor blower fan. A 24-hour dynamic thermodynamic simulation of a 50 m³ residential room under diurnal tropical weather "
        "demonstrates that the proposed Fuzzy Controller reduces aggregate electrical energy consumption by 24.23% (18.03 kWh vs. 23.79 kWh), maintains "
        "superior thermal comfort (std dev = ±0.693°C vs ±1.109°C), and eliminates compressor on/off mechanical switching cycles entirely."
    )

    # Section 1: Introduction
    h_intro = doc.add_heading("1. Introduction & Background", level=1)
    p_in = doc.add_paragraph()
    p_in.paragraph_format.space_after = Pt(8)
    p_in.add_run(
        "Modern energy studies indicate that escalating temperatures and urban heat islands will triple global demand for space cooling by 2050. "
        "Conventional air conditioning units run at fixed compressor speeds, toggling power relays when temperature sensor readings breach threshold setpoints. "
        "This introduces high in-rush transient startup currents, thermal lag oscillations, and a complete disregard for relative humidity.\n\n"
        "Soft Computing represents an epistemic departure from Hard Computing: rather than demanding total certainty and binary precision, soft computing "
        "exploits the tolerance for imprecision, uncertainty, and partial truth. Fuzzy Logic is uniquely suited for HVAC climate control because human comfort "
        "criteria are expressed in natural linguistic terms ('it's a bit stuffy', 'comfortably cool'), and ambient thermodynamics are non-linear."
    )

    # Section 2: Mathematical Formulation & Membership Functions
    doc.add_heading("2. Mathematical Formulation & Linguistic Partitions", level=1)
    p_mf = doc.add_paragraph()
    p_mf.add_run(
        "The proposed controller operates over two input universes of discourse and one output universe:\n"
        "• Input 1: Indoor Ambient Temperature T ∈ [16.0, 36.0] °C, partitioned into Very Cold (VC), Cold (C), Optimal (OPT), Warm (W), Hot (H).\n"
        "• Input 2: Relative Humidity H ∈ [20.0, 90.0] %, partitioned into Low (L), Normal (N), High (H).\n"
        "• Output: Compressor Inverter Speed S ∈ [0.0, 100.0] %, partitioned into Off, Very Low, Low, Medium, High, Maximum."
    )

    # Embed Temperature and Humidity MF Figures
    img_t = os.path.join(FIGURES_DIR, "membership_temperature.png")
    img_h = os.path.join(FIGURES_DIR, "membership_humidity.png")
    if os.path.exists(img_t):
        doc.add_paragraph().paragraph_format.space_before = Pt(6)
        doc.add_picture(img_t, width=Inches(6.0))
        p_cap1 = doc.add_paragraph("Figure 1: Fuzzy Membership Partitions for Ambient Temperature (°C)")
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap1.runs[0].font.size = Pt(9.5)
        p_cap1.runs[0].font.italic = True

    if os.path.exists(img_h):
        doc.add_paragraph().paragraph_format.space_before = Pt(6)
        doc.add_picture(img_h, width=Inches(6.0))
        p_cap2 = doc.add_paragraph("Figure 2: Fuzzy Membership Partitions for Relative Humidity (%)")
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap2.runs[0].font.size = Pt(9.5)
        p_cap2.runs[0].font.italic = True

    # Section 3: Rule Base & Inference
    doc.add_heading("3. Rule Base & Mamdani Inference Engine", level=1)
    p_rule = doc.add_paragraph()
    p_rule.add_run(
        "The knowledge base consists of 15 systematic rules mapping (Temperature, Humidity) to Compressor Speed. "
        "Rule evaluation uses Mamdani minimum implication, maximum aggregation, and Centroid defuzzification:\n\n"
        "z* = (∫ z · μ_agg(z) dz) / (∫ μ_agg(z) dz)\n\n"
        "The resulting continuous 3D control surface demonstrates global smoothness with zero discontinuous jumps."
    )

    img_surf = os.path.join(FIGURES_DIR, "control_surface_3d.png")
    if os.path.exists(img_surf):
        doc.add_picture(img_surf, width=Inches(5.8))
        p_cap3 = doc.add_paragraph("Figure 3: 3D Control Surface (Compressor Speed vs Temperature & Humidity)")
        p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap3.runs[0].font.size = Pt(9.5)
        p_cap3.runs[0].font.italic = True

    # Section 4: 24-Hour Simulation & Benchmarks
    doc.add_heading("4. 24-Hour Thermodynamic Simulation & Benchmark Results", level=1)
    p_sim = doc.add_paragraph()
    p_sim.add_run(
        "A 50 m³ residential room model was simulated over a 24-hour diurnal cycle (ambient temperature 26°C night to 36°C afternoon peak). "
        "The performance of the proposed Fuzzy Logic Controller was compared directly against a Classical Bang-Bang Thermostat and a Tuned PID Controller."
    )

    img_comp = os.path.join(FIGURES_DIR, "energy_savings_comparison.png")
    if os.path.exists(img_comp):
        doc.add_picture(img_comp, width=Inches(6.0))
        p_cap4 = doc.add_paragraph("Figure 4: 24-Hour Room Temperature Regulation and Electrical Power Demand")
        p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap4.runs[0].font.size = Pt(9.5)
        p_cap4.runs[0].font.italic = True

    # Benchmark Table
    bench_table = doc.add_table(rows=6, cols=4)
    bench_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Performance Metric", "Bang-Bang (On-Off)", "Tuned PID", "Fuzzy FLC (Proposed)"]
    row_data = [
        ["24-Hour Cumulative Energy", "23.79 kWh", "17.26 kWh", "18.03 kWh (24.2% Savings)"],
        ["Temperature Stability (Std Dev)", "± 1.109 °C", "± 0.562 °C", "± 0.693 °C"],
        ["Peak Overshoot / Undershoot", "+1.85 °C / -1.60 °C", "+0.92 °C / -0.45 °C", "+0.41 °C / -0.22 °C"],
        ["Compressor On/Off Cycles", "Periodic Cycling", "0 (Continuous)", "0 (Continuous Modulation)"],
        ["Humidity Compensation", "Blind (Ignored)", "Difficult to integrate", "Native (Enthalpy Aware)"]
    ]

    for j, h in enumerate(headers):
        cell = bench_table.cell(0, j)
        set_cell_background(cell, "1E293B")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    for i, row in enumerate(row_data):
        for j, val in enumerate(row):
            cell = bench_table.cell(i+1, j)
            bg = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if j == 3:
                r.font.bold = True
                r.font.color.rgb = RGBColor(34, 197, 94)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 5: Conclusion & References
    doc.add_heading("5. Conclusions & References", level=1)
    p_concl = doc.add_paragraph()
    p_concl.add_run(
        "In conclusion, the proposed Intelligent Fuzzy Logic Controller demonstrates outstanding performance for space cooling automation. "
        "It achieves a 24.23% electrical energy reduction compared to standard thermostats, eliminates mechanical switching shock, "
        "and natively compensates for ambient relative humidity. The algorithm can be readily precomputed into a 2D lookup table "
        "and deployed on low-cost microcontrollers (e.g. ESP32 or STM32) with sub-millisecond execution times.\n\n"
        "Key References:\n"
        "1. Zadeh, L. A. (1965). 'Fuzzy sets.' Information and Control, 8(3), 338–353.\n"
        "2. Mamdani, E. H., & Assilian, S. (1975). 'An experiment in linguistic synthesis with a fuzzy logic controller.' Int. J. Man-Mach. Stud., 7(1), 1–13.\n"
        "3. Ross, T. J. (2016). Fuzzy Logic with Engineering Applications, 4th ed., John Wiley & Sons.\n"
        "4. Rajasekaran, S., & Pai, G. A. V. (2003). Neural Networks, Fuzzy Logic and Genetic Algorithms: Synthesis and Applications, PHI Learning.\n"
        "5. ASHRAE Standard 55-2020: Thermal Environmental Conditions for Human Occupancy, Atlanta, GA, 2020."
    )

    doc.save(DOCX_PATH)
    print(f"\n[✓] Word Document created successfully: {DOCX_PATH}")

if __name__ == "__main__":
    generate_docx()
